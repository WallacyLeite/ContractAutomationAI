"""
Script utilitário — Cria o template Word base do contrato.

Execute uma única vez para gerar o arquivo:
    python create_template.py

O template usa a sintaxe do docxtpl: {{ variavel }}
As variáveis são substituídas automaticamente pelo backend ao gerar cada contrato.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def criar_template():
    doc = Document()

    # -----------------------------------------------------------------------
    # Configuração de margem da página
    # -----------------------------------------------------------------------
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # -----------------------------------------------------------------------
    # Título principal
    # -----------------------------------------------------------------------
    titulo = doc.add_heading("CONTRATO DE PRESTAÇÃO DE SERVIÇOS", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].font.size = Pt(16)
    titulo.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Seção: Dados do Contratante
    # -----------------------------------------------------------------------
    doc.add_heading("1. PARTES CONTRATANTES", level=2)

    def linha(doc, rotulo, variavel):
        p = doc.add_paragraph()
        run_label = p.add_run(f"{rotulo}: ")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_var = p.add_run(variavel)
        run_var.font.size = Pt(11)
        return p

    linha(doc, "CONTRATANTE", "{{ nome_cliente }}")
    linha(doc, "ENDEREÇO", "{{ endereco_cliente }}")
    linha(doc, "CPF / CNPJ", "{{ documento_cliente }}")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Seção: Objeto do Contrato
    # -----------------------------------------------------------------------
    doc.add_heading("2. OBJETO DO CONTRATO", level=2)
    linha(doc, "Tipo de Serviço", "{{ tipo_servico }}")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Seção: Valor e Prazo
    # -----------------------------------------------------------------------
    doc.add_heading("3. VALOR E PRAZO", level=2)
    linha(doc, "Valor Total", "{{ valor }}")
    linha(doc, "Data de Início", "{{ data_inicio }}")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Seção: Observações
    # -----------------------------------------------------------------------
    doc.add_heading("4. OBSERVAÇÕES E CONDIÇÕES ESPECIAIS", level=2)
    p_obs = doc.add_paragraph("{{ observacoes }}")
    p_obs.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Assinaturas
    # -----------------------------------------------------------------------
    doc.add_heading("5. ASSINATURAS", level=2)

    doc.add_paragraph()

    p_linha = doc.add_paragraph("_" * 45 + "          " + "_" * 45)
    p_linha.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_labels = doc.add_paragraph("CONTRATANTE" + " " * 48 + "CONTRATADA")
    p_labels.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_labels.runs:
        run.bold = True
        run.font.size = Pt(11)

    doc.add_paragraph()

    p_local = doc.add_paragraph("Local e Data: _____________________________________________")
    p_local.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -----------------------------------------------------------------------
    # Salva o template na pasta correta
    # -----------------------------------------------------------------------
    os.makedirs("templates", exist_ok=True)
    caminho = os.path.join("templates", "contrato_modelo.docx")
    doc.save(caminho)
    print(f"Template criado com sucesso: {caminho}")
    print("\nVariáveis disponíveis no template:")
    variaveis = [
        "{{ nome_cliente }}",
        "{{ endereco_cliente }}",
        "{{ documento_cliente }}",
        "{{ tipo_servico }}",
        "{{ valor }}",
        "{{ data_inicio }}",
        "{{ observacoes }}",
    ]
    for v in variaveis:
        print(f"  {v}")


if __name__ == "__main__":
    criar_template()
